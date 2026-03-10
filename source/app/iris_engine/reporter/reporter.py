#  IRIS Source Code
#  Copyright (C) 2022 - DFIR IRIS Team
#  contact@dfir-iris.org
#  Copyright (C) 2021 - Airbus CyberSecurity (SAS)
#  ir@cyberactionlab.net
#
#  This program is free software; you can redistribute it and/or
#  modify it under the terms of the GNU Lesser General Public
#  License as published by the Free Software Foundation; either
#  version 3 of the License, or (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the GNU
#  Lesser General Public License for more details.
#
#  You should have received a copy of the GNU Lesser General Public License
#  along with this program; if not, write to the Free Software Foundation,
#  Inc., 51 Franklin Street, Fifth Floor, Boston, MA  02110-1301, USA.

# IMPORTS ------------------------------------------------

# VARS ---------------------------------------------------

# CONTENT ------------------------------------------------
import logging as log
import os
import uuid
from datetime import datetime

import jinja2
from jinja2.sandbox import SandboxedEnvironment

from app.datamgmt.reporter.report_db import export_case_json_for_report
from app.iris_engine.utils.common import IrisJinjaEnv
from docx_generator.docx_generator import DocxGenerator
from docx_generator.exceptions import rendering_error
from flask_login import current_user
from sqlalchemy import desc

from app import app
from app.datamgmt.activities.activities_db import get_auto_activities
from app.datamgmt.activities.activities_db import get_manual_activities
from app.datamgmt.case.case_db import case_get_desc_crc
from app.datamgmt.reporter.report_db import export_case_json
from app.models import AssetsType
from app.models import CaseAssets
from app.models import CaseEventsAssets
from app.models import CaseReceivedFile
from app.models import CaseTemplateReport
from app.models import CasesEvent
from app.models import Ioc
from app.models import IocAssetLink
from app.models import IocLink
from app.iris_engine.reporter.ImageHandler import ImageHandler

LOG_FORMAT = '%(asctime)s :: %(levelname)s :: %(module)s :: %(funcName)s :: %(message)s'
log.basicConfig(level=log.INFO, format=LOG_FORMAT)


def _inject_timeline_image_into_docx(docx_path, image_path):
    """
    Post-processes the generated DOCX to insert the visual timeline PNG
    after the TIMELINE section heading, using python-docx directly.
    This bypasses the autoescape issue with docxtpl's addPicture.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        doc = Document(docx_path)

        # Find the paragraph/table that corresponds to the TIMELINE annex
        # Strategy: look for the last table in the document (the timeline table)
        # and insert the image paragraph right after it.
        tables = doc.tables
        if not tables:
            return

        # Find the timeline table — it's the one with 'Date' / 'Event' columns
        # Fall back to the last table in the document
        timeline_table = None
        for tbl in tables:
            if tbl.rows:
                header_texts = ' '.join(
                    cell.text.strip().lower() for cell in tbl.rows[0].cells
                )
                if 'date' in header_texts and ('event' in header_texts or 'timeline' in header_texts):
                    timeline_table = tbl
                    break
        if timeline_table is None:
            timeline_table = tables[-1]

        # Insert a new paragraph after the table using raw XML manipulation
        tbl_element = timeline_table._tbl
        tbl_parent = tbl_element.getparent()

        # Create a new paragraph with centered alignment
        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        new_para.append(pPr)

        # Add a run with the picture to the new paragraph
        run_elem = OxmlElement('w:r')
        new_para.append(run_elem)

        # Insert the paragraph after the table in the XML tree
        tbl_element.addnext(new_para)

        # Now use python-docx to add the picture to that run by finding the
        # new paragraph in doc.paragraphs and using add_run().add_picture()
        # Find the new paragraph object by its XML element
        para_obj = None
        for p in doc.paragraphs:
            if p._p is new_para:
                para_obj = p
                break

        if para_obj is None:
            # Fallback: append to document body
            para_obj = doc.add_paragraph()
            para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Get available page width minus margins
        section = doc.sections[0]
        page_w = section.page_width - section.left_margin - section.right_margin

        run_obj = para_obj.add_run()
        pic = run_obj.add_picture(image_path)
        # Scale to full page width
        if pic.width > page_w:
            ratio = pic.height / pic.width
            pic.width = page_w
            pic.height = int(ratio * page_w)
        para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(docx_path)
        log.info('Timeline image injected into report: %s', docx_path)

    except Exception as exc:
        log.warning('Failed to inject timeline image into DOCX: %s', exc)


def _generate_timeline_image(timeline_events, output_dir):
    """
    Generates a visual timeline PNG image from case events using Pillow.
    Design: card-based with colored left accent bars, day-group badges and a
    connecting spine — mirroring the IRIS web UI timeline style.
    Returns the path to the generated image, or None on error / no events.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        from collections import OrderedDict

        FONT_BOLD = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'
        FONT      = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'

        # ── layout ────────────────────────────────────────────────────────
        W        = 1800
        HEADER_H = 78
        FOOTER_H = 46
        SPINE_X  = 72          # absolute X of the vertical spine
        CARD_X   = SPINE_X + 46
        CARD_W   = W - CARD_X - 40
        ACCENT_W = 8           # colored left edge of card
        CARD_PAD = 20          # vertical gap between cards
        BADGE_H  = 42
        BADGE_MB = 18
        DOT_R    = 11          # radius of spine dot
        TEXT_PAD = ACCENT_W + 22
        MAX_CONTENT_LINES = 3
        LINE_H_CONTENT    = 23

        # ── colors ────────────────────────────────────────────────────────
        C_BG      = '#f5f6fa'
        C_CARD    = '#ffffff'
        C_BORDER  = '#dee2e6'
        C_SHADOW  = '#cdd1d9'
        C_HEADER  = '#1e2a3a'
        C_HEAD_FG = '#ffffff'
        C_BADGE   = '#343a40'
        C_BADGE_FG= '#ffffff'
        C_SPINE   = '#adb5bd'
        C_TITLE   = '#212529'
        C_CONTENT = '#495057'
        C_FOOTER  = '#868e96'
        C_DEFAULT = '#4361ee'

        # ── fonts ─────────────────────────────────────────────────────────
        try:
            f_head   = ImageFont.truetype(FONT_BOLD, 34)
            f_badge  = ImageFont.truetype(FONT_BOLD, 22)
            f_title  = ImageFont.truetype(FONT_BOLD, 24)
            f_time   = ImageFont.truetype(FONT_BOLD, 18)
            f_body   = ImageFont.truetype(FONT, 18)
            f_footer = ImageFont.truetype(FONT, 16)
        except Exception:
            f_head = f_badge = f_title = f_time = f_body = f_footer = \
                ImageFont.load_default()

        # helper: pixel width of text with a given font
        def _tw(font, text):
            try:
                return int(font.getlength(text))
            except AttributeError:
                return font.getsize(text)[0]

        # helper: word-wrap text to fit inside max_px pixels
        def _wrap(text, font, max_px):
            words = text.split()
            lines, curr = [], []
            for w in words:
                probe = ' '.join(curr + [w])
                if _tw(font, probe) <= max_px:
                    curr.append(w)
                else:
                    if curr:
                        lines.append(' '.join(curr))
                    curr = [w]
                    if len(lines) >= MAX_CONTENT_LINES:
                        break
            if curr and len(lines) < MAX_CONTENT_LINES:
                lines.append(' '.join(curr))
            return lines

        # ── sort & group by day ───────────────────────────────────────────
        events = sorted(
            [e for e in timeline_events if e.get('event_date')],
            key=lambda e: e['event_date']
        )
        if not events:
            return None

        groups = OrderedDict()
        for ev in events:
            d = ev['event_date']
            k = d.strftime('%Y-%m-%d') if hasattr(d, 'strftime') else str(d)[:10]
            groups.setdefault(k, []).append(ev)

        # ── precompute content lines & card heights ───────────────────────
        text_area_w = CARD_W - TEXT_PAD - 20
        ev_content_lines = {}
        ev_card_h = {}
        for ev in events:
            raw = str(ev.get('event_content') or '').replace('\n', ' ').strip()
            lines = _wrap(raw, f_body, text_area_w) if raw else []
            ev_content_lines[id(ev)] = lines
            ch = 18 + 30 + (len(lines) * LINE_H_CONTENT if lines else 0) + 20
            ev_card_h[id(ev)] = max(86, ch)

        # ── compute canvas height ─────────────────────────────────────────
        total_h = HEADER_H + CARD_PAD
        for day_key, day_events in groups.items():
            total_h += BADGE_H + BADGE_MB
            for ev in day_events:
                total_h += ev_card_h[id(ev)] + CARD_PAD
        total_h += FOOTER_H

        # ── render ────────────────────────────────────────────────────────
        img  = Image.new('RGB', (W, total_h), color=C_BG)
        draw = ImageDraw.Draw(img)

        # header bar (gradient simulated with two horizontal bands)
        draw.rectangle([(0, 0), (W, HEADER_H)], fill=C_HEADER)
        draw.rectangle([(0, HEADER_H - 4), (W, HEADER_H)], fill='#253550')
        draw.text((W // 2, HEADER_H // 2), 'Case Timeline',
                  fill=C_HEAD_FG, font=f_head, anchor='mm')

        # vertical spine
        draw.line([(SPINE_X, HEADER_H), (SPINE_X, total_h - FOOTER_H)],
                  fill=C_SPINE, width=2)

        y = HEADER_H + CARD_PAD
        for day_key, day_events in groups.items():

            # day badge
            bw = _tw(f_badge, day_key) + 32
            bx1 = SPINE_X - 12
            draw.rectangle([(bx1, y), (bx1 + bw, y + BADGE_H)], fill=C_BADGE)
            draw.text((bx1 + 16, y + BADGE_H // 2), day_key,
                      fill=C_BADGE_FG, font=f_badge, anchor='lm')
            y += BADGE_H + BADGE_MB

            for ev in day_events:
                ch     = ev_card_h[id(ev)]
                cy     = y
                cy_mid = cy + ch // 2

                raw_color = ev.get('event_color') or C_DEFAULT
                if not isinstance(raw_color, str) \
                        or not raw_color.startswith('#') \
                        or len(raw_color) not in (4, 7):
                    raw_color = C_DEFAULT

                # card shadow
                draw.rectangle(
                    [(CARD_X + 4, cy + 4), (CARD_X + CARD_W + 4, cy + ch + 4)],
                    fill=C_SHADOW
                )
                # card body
                draw.rectangle(
                    [(CARD_X, cy), (CARD_X + CARD_W, cy + ch)],
                    fill=C_CARD, outline=C_BORDER, width=1
                )
                # colored left accent bar (like web UI border-left)
                draw.rectangle(
                    [(CARD_X, cy), (CARD_X + ACCENT_W, cy + ch)],
                    fill=raw_color
                )

                # spine dot
                draw.ellipse(
                    [(SPINE_X - DOT_R, cy_mid - DOT_R),
                     (SPINE_X + DOT_R, cy_mid + DOT_R)],
                    fill=raw_color, outline=C_BG, width=2
                )
                # horizontal connector: dot → card
                draw.line(
                    [(SPINE_X + DOT_R, cy_mid), (CARD_X, cy_mid)],
                    fill=raw_color, width=2
                )

                # ── card text ─────────────────────────────────────────────
                tx = CARD_X + TEXT_PAD
                ty = cy + 16

                # time pill
                ev_date = ev.get('event_date')
                if hasattr(ev_date, 'strftime'):
                    time_str = ev_date.strftime('%H:%M')
                else:
                    time_str = str(ev_date)[11:16] if ev_date else '--:--'
                pill_w = _tw(f_time, time_str) + 20
                pill_h = 28
                draw.rectangle(
                    [(tx, ty), (tx + pill_w, ty + pill_h)],
                    fill=raw_color
                )
                draw.text((tx + 10, ty + pill_h // 2), time_str,
                          fill='#ffffff', font=f_time, anchor='lm')

                # event title (right of pill)
                title = str(ev.get('event_title') or '')[:110]
                draw.text((tx + pill_w + 18, ty + 2), title,
                          fill=C_TITLE, font=f_title)

                # content lines
                for li, line in enumerate(ev_content_lines[id(ev)]):
                    draw.text(
                        (tx, ty + pill_h + 10 + li * LINE_H_CONTENT),
                        line, fill=C_CONTENT, font=f_body
                    )

                y += ch + CARD_PAD

        # footer
        draw.text((W // 2, total_h - FOOTER_H // 2),
                  'Generated by IRIS',
                  fill=C_FOOTER, font=f_footer, anchor='mm')

        out_path = os.path.join(
            output_dir, 'timeline_{}.png'.format(uuid.uuid4().hex)
        )
        img.save(out_path, 'PNG', dpi=(150, 150))
        return out_path

    except Exception as exc:
        log.warning('Timeline image generation failed: %s', exc)
        return None


class IrisReportMaker(object):
    """
    IRIS generical report maker
    """

    def __init__(self, tmp_dir, report_id, caseid, safe_mode=False):
        self._tmp = tmp_dir
        self._report_id = report_id
        self._case_info = {}
        self._caseid = caseid
        self.safe_mode = safe_mode

    def get_case_info(self, doc_type):
        """Returns case information

        Args:
            doc_type (_type_): Investigation or Activities report

        Returns:
            _type_: case info
        """
        if doc_type == 'Investigation':
            case_info = self._get_case_info()
        elif doc_type == 'Activities':
            case_info = self._get_activity_info()
        else:
            log.error("Unknown report type")
            return None
        return case_info

    def _get_activity_info(self):
        auto_activities = get_auto_activities(self._caseid)
        manual_activities = get_manual_activities(self._caseid)
        case_info_in = self._get_case_info()

        # Format information and generate the activity report #
        doc_id = "{}".format(datetime.utcnow().strftime("%y%m%d_%H%M"))

        case_info = {
            'auto_activities': auto_activities,
            'manual_activities': manual_activities,
            'date': datetime.utcnow(),
            'gen_user': current_user.name,
            'case': {'name': case_info_in['case'].get('name'),
                     'open_date': case_info_in['case'].get('open_date'),
                     'for_customer': case_info_in['case'].get('client').get('customer_name'),
                     'client': case_info_in['case'].get('client')
                     },
            'doc_id': doc_id
        }

        return case_info

    def _get_case_info(self):
        """
        Retrieve information of the case
        :return:
        """
        case_info = export_case_json(self._caseid)

        # Get customer, user and case title
        case_info['doc_id'] = IrisReportMaker.get_docid()
        case_info['user'] = current_user.name

        # Set date
        case_info['date'] = datetime.utcnow().strftime("%Y-%m-%d")

        return case_info

    @staticmethod
    def get_case_summary(caseid):
        """
        Retrieve the case summary from thehive
        :return:
        """

        _crc32, descr = case_get_desc_crc(caseid)

        # return IrisMakeDocReport.markdown_to_text(descr)
        return descr

    @staticmethod
    def get_case_files(caseid):
        """
        Retrieve the list of files with their hashes
        :return:
        """
        files = CaseReceivedFile.query.filter(
            CaseReceivedFile.case_id == caseid
        ).with_entities(
            CaseReceivedFile.filename,
            CaseReceivedFile.date_added,
            CaseReceivedFile.file_hash,
            CaseReceivedFile.custom_attributes
        ).order_by(
            CaseReceivedFile.date_added
        ).all()

        if files:
            return [row._asdict() for row in files]

        else:
            return []

    @staticmethod
    def get_case_timeline(caseid):
        """
        Retrieve the case timeline
        :return:
        """
        timeline = CasesEvent.query.filter(
            CasesEvent.case_id == caseid
        ).order_by(
            CasesEvent.event_date
        ).all()

        cache_id = {}
        ras = {}
        tim = []
        for row in timeline:
            ras = row
            setattr(ras, 'asset', None)

            as_list = CaseEventsAssets.query.with_entities(
                CaseAssets.asset_id,
                CaseAssets.asset_name,
                AssetsType.asset_name.label('type')
            ).filter(
                CaseEventsAssets.event_id == row.event_id
            ).join(CaseEventsAssets.asset, CaseAssets.asset_type).all()

            alki = []
            for asset in as_list:
                alki.append("{} ({})".format(asset.asset_name, asset.type))

            setattr(ras, 'asset', "\r\n".join(alki))

            tim.append(ras)

        return tim

    @staticmethod
    def get_case_ioc(caseid):
        """
        Retrieve the list of IOC linked to the case
        :return:
        """
        res = IocLink.query.distinct().with_entities(
            Ioc.ioc_value,
            Ioc.ioc_type,
            Ioc.ioc_description,
            Ioc.ioc_tags,
            Ioc.custom_attributes
        ).filter(
            IocLink.case_id == caseid
        ).join(IocLink.ioc).order_by(Ioc.ioc_type).all()

        if res:
            return [row._asdict() for row in res]

        else:
            return []

    @staticmethod
    def get_case_assets(caseid):
        """
        Retrieve the assets linked ot the case
        :return:
        """
        ret = []

        res = CaseAssets.query.distinct().with_entities(
            CaseAssets.asset_id,
            CaseAssets.asset_name,
            CaseAssets.asset_description,
            CaseAssets.asset_compromised.label('compromised'),
            AssetsType.asset_name.label("type"),
            CaseAssets.custom_attributes,
            CaseAssets.asset_tags
        ).filter(
            CaseAssets.case_id == caseid
        ).join(
            CaseAssets.asset_type
        ).order_by(desc(CaseAssets.asset_compromised)).all()

        for row in res:
            row = row._asdict()
            row['light_asset_description'] = row['asset_description']

            ial = IocAssetLink.query.with_entities(
                Ioc.ioc_value,
                Ioc.ioc_type,
                Ioc.ioc_description
            ).filter(
                IocAssetLink.asset_id == row['asset_id']
            ).join(
                IocAssetLink.ioc
            ).all()

            if ial:
                row['asset_ioc'] = [row._asdict() for row in ial]
            else:
                row['asset_ioc'] = []

            ret.append(row)

        return ret

    @staticmethod
    def get_docid():
        return "{}".format(
            datetime.utcnow().strftime("%y%m%d_%H%M"))

    @staticmethod
    def markdown_to_text(markdown_string):
        """
        Converts a markdown string to plaintext
        """
        return markdown_string.replace('\n', '</w:t></w:r><w:r/></w:p><w:p><w:r><w:t xml:space="preserve">').replace(
            '#', '')


class IrisMakeDocReport(IrisReportMaker):
    """
    Generates a DOCX report for the case
    """

    def __init__(self, tmp_dir, report_id, caseid, safe_mode=False):
        self._tmp = tmp_dir
        self._report_id = report_id
        self._case_info = {}
        self._caseid = caseid
        self._safe_mode = safe_mode

    def generate_doc_report(self, doc_type):
        """
        Actually generates the report
        :return:
        """
        if doc_type == 'Investigation':
            case_info = self._get_case_info()
        elif doc_type == 'Activities':
            case_info = self._get_activity_info()
        else:
            log.error("Unknown report type")
            return None

        report = CaseTemplateReport.query.filter(CaseTemplateReport.id == self._report_id).first()

        name = "{}".format("{}.docx".format(report.naming_format))
        name = name.replace("%code_name%", case_info['doc_id'])
        name = name.replace('%customer%', case_info['case']['client']['customer_name'])
        name = name.replace('%case_name%', case_info['case'].get('name'))
        name = name.replace('%date%', datetime.utcnow().strftime("%Y-%m-%d"))
        output_file_path = os.path.join(self._tmp, name)

        try:

            if not self._safe_mode:
                image_handler = ImageHandler(template=None, base_path='/')
            else:
                image_handler = None

            generator = DocxGenerator(image_handler=image_handler)
            generator.generate_docx("/",
                                    os.path.join(app.config['TEMPLATES_PATH'], report.internal_reference),
                                    case_info,
                                    output_file_path
                                    )

            # Post-process: inject visual timeline image into the generated DOCX
            timeline_img = case_info.get('timeline_image', '')
            if timeline_img and os.path.isfile(timeline_img) and doc_type == 'Investigation':
                _inject_timeline_image_into_docx(output_file_path, timeline_img)

            return output_file_path, ""

        except rendering_error.RenderingError as e:

            return None, e.__str__()

    def _get_activity_info(self):
        auto_activities = get_auto_activities(self._caseid)
        manual_activities = get_manual_activities(self._caseid)
        case_info_in = self._get_case_info()

        # Format information and generate the activity report #
        doc_id = "{}".format(datetime.utcnow().strftime("%y%m%d_%H%M"))

        case_info = {
            'auto_activities': auto_activities,
            'manual_activities': manual_activities,
            'date': datetime.utcnow(),
            'gen_user': current_user.name,
            'case': {'name': case_info_in['case'].get('name'),
                     'open_date': case_info_in['case'].get('open_date'),
                     'for_customer': case_info_in['case'].get('for_customer'),
                     'client': case_info_in['case'].get('client')
                     },
            'doc_id': doc_id
        }

        return case_info

    def _get_case_info(self):
        """
        Retrieve information of the case
        :return:
        """
        case_info = export_case_json_for_report(self._caseid)

        # Get customer, user and case title
        case_info['doc_id'] = IrisMakeDocReport.get_docid()
        case_info['user'] = current_user.name

        # Set date
        case_info['date'] = datetime.utcnow().strftime("%Y-%m-%d")

        # Generate visual timeline image and expose its path to the template
        timeline_events = case_info.get('timeline', [])
        img_path = _generate_timeline_image(timeline_events, self._tmp)
        case_info['timeline_image'] = img_path or ''

        return case_info

    @staticmethod
    def get_case_summary(caseid):
        """
        Retrieve the case summary from thehive
        :return:
        """

        _crc32, descr = case_get_desc_crc(caseid)

        # return IrisMakeDocReport.markdown_to_text(descr)
        return descr

    @staticmethod
    def get_case_files(caseid):
        """
        Retrieve the list of files with their hashes
        :return:
        """
        files = CaseReceivedFile.query.filter(
            CaseReceivedFile.case_id == caseid
        ).with_entities(
            CaseReceivedFile.filename,
            CaseReceivedFile.date_added,
            CaseReceivedFile.file_hash,
            CaseReceivedFile.custom_attributes
        ).order_by(
            CaseReceivedFile.date_added
        ).all()

        if files:
            return [row._asdict() for row in files]

        else:
            return []

    @staticmethod
    def get_case_timeline(caseid):
        """
        Retrieve the case timeline
        :return:
        """
        timeline = CasesEvent.query.filter(
            CasesEvent.case_id == caseid
        ).order_by(
            CasesEvent.event_date
        ).all()

        cache_id = {}
        ras = {}
        tim = []
        for row in timeline:
            ras = row
            setattr(ras, 'asset', None)

            as_list = CaseEventsAssets.query.with_entities(
                CaseAssets.asset_id,
                CaseAssets.asset_name,
                AssetsType.asset_name.label('type')
            ).filter(
                CaseEventsAssets.event_id == row.event_id
            ).join(CaseEventsAssets.asset, CaseAssets.asset_type).all()

            alki = []
            for asset in as_list:
                alki.append("{} ({})".format(asset.asset_name, asset.type))

            setattr(ras, 'asset', "\r\n".join(alki))

            tim.append(ras)

        return tim

    @staticmethod
    def get_case_ioc(caseid):
        """
        Retrieve the list of IOC linked to the case
        :return:
        """
        res = IocLink.query.distinct().with_entities(
            Ioc.ioc_value,
            Ioc.ioc_type,
            Ioc.ioc_description,
            Ioc.ioc_tags,
            Ioc.custom_attributes
        ).filter(
            IocLink.case_id == caseid
        ).join(IocLink.ioc).order_by(Ioc.ioc_type).all()

        if res:
            return [row._asdict() for row in res]

        else:
            return []

    @staticmethod
    def get_case_assets(caseid):
        """
        Retrieve the assets linked ot the case
        :return:
        """
        ret = []

        res = CaseAssets.query.distinct().with_entities(
            CaseAssets.asset_id,
            CaseAssets.asset_name,
            CaseAssets.asset_description,
            CaseAssets.asset_compromise_status_id.label('compromise_status'),
            AssetsType.asset_name.label("type"),
            CaseAssets.custom_attributes,
            CaseAssets.asset_tags
        ).filter(
            CaseAssets.case_id == caseid
        ).join(
            CaseAssets.asset_type
        ).order_by(desc(CaseAssets.asset_compromise_status_id)).all()

        for row in res:
            row = row._asdict()
            row['light_asset_description'] = row['asset_description']

            ial = IocAssetLink.query.with_entities(
                Ioc.ioc_value,
                Ioc.ioc_type,
                Ioc.ioc_description
            ).filter(
                IocAssetLink.asset_id == row['asset_id']
            ).join(
                IocAssetLink.ioc
            ).all()

            if ial:
                row['asset_ioc'] = [row._asdict() for row in ial]
            else:
                row['asset_ioc'] = []

            ret.append(row)

        return ret

    @staticmethod
    def get_docid():
        return "{}".format(
            datetime.utcnow().strftime("%y%m%d_%H%M"))

    @staticmethod
    def markdown_to_text(markdown_string):
        """
        Converts a markdown string to plaintext
        """
        return markdown_string.replace('\n', '</w:t></w:r><w:r/></w:p><w:p><w:r><w:t xml:space="preserve">').replace(
            '#', '')


class IrisMakeMdReport(IrisReportMaker):
    """
    Generates a MD report for the case
    """

    def __init__(self, tmp_dir, report_id, caseid, safe_mode=False):
        self._tmp = tmp_dir
        self._report_id = report_id
        self._case_info = {}
        self._caseid = caseid
        self.safe_mode = safe_mode

    def generate_md_report(self, doc_type):
        """
        Generate report file
        """
        case_info = self.get_case_info(doc_type)
        if case_info is None:
            return None

        # Get file extension
        report = CaseTemplateReport.query.filter(
            CaseTemplateReport.id == self._report_id).first()

        _, report_format = os.path.splitext(report.internal_reference)

        case_info['case']['for_customer'] = f"{case_info['case'].get('client').get('customer_name')} (legacy::use client.customer_name)"

        # Prepare report name
        name = "{}".format(("{}" + str(report_format)).format(report.naming_format))
        name = name.replace("%code_name%", case_info['doc_id'])
        name = name.replace(
            '%customer%', case_info['case'].get('client').get('customer_name'))
        name = name.replace('%case_name%', case_info['case'].get('name'))
        name = name.replace('%date%', datetime.utcnow().strftime("%Y-%m-%d"))

        # Build output file
        output_file_path = os.path.join(self._tmp, name)

        try:
            env = IrisJinjaEnv()
            env.filters = app.jinja_env.filters
            template = env.from_string(
               open(os.path.join(app.config['TEMPLATES_PATH'], report.internal_reference)).read())
            output_text = template.render(case_info)

            # Write the result in the output file
            with open(output_file_path, 'w', encoding="utf-8") as html_file:
                html_file.write(output_text)

        except Exception as e:
            log.exception("Error while generating report: {}".format(e))
            return None, e.__str__()

        return output_file_path, 'Report generated'


class QueuingHandler(log.Handler):
    """A thread safe logging.Handler that writes messages into a queue object.

       Designed to work with LoggingWidget so log messages from multiple
       threads can be shown together in a single ttk.Frame.

       The standard logging.QueueHandler/logging.QueueListener can not be used
       for this because the QueueListener runs in a private thread, not the
       main thread.

       Warning:  If multiple threads are writing into this Handler, all threads
       must be joined before calling logging.shutdown() or any other log
       destinations will be corrupted.
    """

    def __init__(self, *args, task_self, message_queue, **kwargs):
        """Initialize by copying the queue and sending everything else to superclass."""
        log.Handler.__init__(self, *args, **kwargs)
        self.message_queue = message_queue
        self.task_self = task_self

    def emit(self, record):
        """Add the formatted log message (sans newlines) to the queue."""
        self.message_queue.append(self.format(record).rstrip('\n'))
        self.task_self.update_state(state='PROGRESS',
                                    meta=list(self.message_queue))
