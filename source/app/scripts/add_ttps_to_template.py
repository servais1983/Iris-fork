"""
Script to add a MITRE ATT&CK TTPs section to the IRIS DOCX report template.
Run inside the iriswebapp_app container or locally with python-docx installed.

Usage:
  python add_ttps_to_template.py
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

TEMPLATE_PATH = '/iriswebapp/app/templates/docx_reports/iris_report_template.docx'


def _clear_paragraph(para):
    """Remove all runs from a paragraph, keeping paragraph XML."""
    for child in list(para._element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            para._element.remove(child)


def _set_paragraph_text(para, text):
    """Set the text of a paragraph by replacing all runs with one run."""
    _clear_paragraph(para)
    new_run = OxmlElement('w:r')
    r_pr = para._element.find(qn('w:rPr'))
    if r_pr is not None:
        new_run.append(deepcopy(r_pr))
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_run.append(t)
    para._element.append(new_run)


def _make_paragraph_element(style_id=None, text=''):
    """Create a new <w:p> element with optional style and text."""
    p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
    return p


def _make_table_cell(text, is_header=False):
    """Create a <w:tc> element with text."""
    tc = OxmlElement('w:tc')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    if is_header:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' ') or text.startswith('{'):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc


def _make_table_row(cells_text, is_header=False):
    """Create a <w:tr> element with given cell texts."""
    tr = OxmlElement('w:tr')
    for text in cells_text:
        tc = _make_table_cell(text, is_header=is_header)
        tr.append(tc)
    return tr


def _copy_table_style(source_table):
    """Get the tblPr (table properties) XML from a source table to reuse."""
    tbl_pr = source_table._tbl.find(qn('w:tblPr'))
    if tbl_pr is not None:
        return deepcopy(tbl_pr)
    return None


def add_ttps_section(doc, reference_ioc_table):
    """
    Add a MITRE ATT&CK TTPs section after the IOCs table.

    Template variables used:
      - ttps: list of dicts with keys:
          mitre_id, name, tactics (list), note, description, url
    """
    body = doc.element.body

    # Find the last {% endif %} paragraph in the document
    # (the one that closes the {% if iocs|count %} block)
    last_endif_para = None
    for p in body.findall(qn('w:p')):
        for r in p.findall('.//' + qn('w:t')):
            if r.text and '{% endif %}' in r.text:
                last_endif_para = p

    if last_endif_para is None:
        print("ERROR: Could not find the closing {% endif %} paragraph.")
        return False

    # We'll insert after last_endif_para
    insert_after = last_endif_para

    def insert_next(ref_elem, new_elem):
        ref_elem.addnext(new_elem)
        return new_elem

    # 1. {% if ttps|count %}
    p_if = _make_paragraph_element(text='{% if ttps|count %}')
    insert_next(insert_after, p_if)
    insert_after = p_if

    # 2. Heading: "MITRE ATT&CK TTPs"
    #    Use the same style as other annexe headings: 'Titre Annexe 1'
    #    Get the style ID from the document styles
    heading_style_id = 'TitreAnnexe1'  # tentative
    # Check if it exists
    styles = doc.styles
    found_style = None
    for s in styles:
        if 'Titre Annexe 1' in s.name or 'TitreAnnexe1' in s.style_id:
            found_style = s.style_id
            break
        if 'annexe' in s.name.lower():
            found_style = s.style_id

    p_heading = _make_paragraph_element(style_id=found_style, text='MITRE ATT\u0026CK TTPs')
    insert_next(insert_after, p_heading)
    insert_after = p_heading

    # 3. Build the TTPs table (4 rows: header, for-loop, data, endfor)
    #    Columns: MITRE ID | Technique | Tactics | Notes
    tbl = OxmlElement('w:tbl')

    # Copy table properties from the IOC table (styling)
    tbl_pr = _copy_table_style(reference_ioc_table)
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
    tbl.append(tbl_pr)

    # Copy tblGrid (column widths) from IOC table - it has 3 cols, we want 4
    # Build our own tblGrid for 4 equally-spaced columns
    tbl_grid = OxmlElement('w:tblGrid')
    for _ in range(4):
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), '2268')  # approx 9072 total / 4
        tbl_grid.append(grid_col)
    tbl.append(tbl_grid)

    # Row 0: Headers
    header_row = _make_table_row(
        ['MITRE ID', 'Technique', 'Tactics', 'Notes'],
        is_header=True
    )
    tbl.append(header_row)

    # Row 1: {%tr for item in ttps %} (docxtpl loop start - same in all cells)
    loop_start_row = _make_table_row(
        ['{%tr for item in ttps %}'] * 4
    )
    tbl.append(loop_start_row)

    # Row 2: Data row
    data_row = _make_table_row([
        '{{ item.mitre_id }}',
        '{{ item.name }}',
        "{{ item.tactics|join(', ') }}",
        '{{ item.note }}'
    ])
    tbl.append(data_row)

    # Row 3: {%tr endfor %}
    loop_end_row = _make_table_row(
        ['{%tr endfor %}'] * 4
    )
    tbl.append(loop_end_row)

    insert_next(insert_after, tbl)
    insert_after = tbl

    # 4. {% endif %}
    p_endif = _make_paragraph_element(text='{% endif %}')
    insert_next(insert_after, p_endif)

    print("TTPs section added successfully.")
    return True


def main():
    print(f"Loading template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    # The IOCs table is the last table in the document (Table 7)
    ioc_table = doc.tables[-1]
    print(f"Reference IOC table: {len(ioc_table.rows)} rows x {len(ioc_table.columns)} cols")
    print(f"IOC table row 0: {[c.text for c in ioc_table.rows[0].cells]}")

    # Check current style IDs
    print("Styles containing 'annexe' or 'Annexe':")
    for s in doc.styles:
        if 'annexe' in s.name.lower() or 'Annexe' in s.name:
            print(f"  name={s.name!r}, id={s.style_id!r}")

    success = add_ttps_section(doc, ioc_table)
    if not success:
        print("FAILED to add TTPs section.")
        return

    # Verify new table count
    print(f"Tables after modification: {len(doc.tables)}")
    last_table = doc.tables[-1]
    print(f"Last table rows: {len(last_table.rows)}, cols: {len(last_table.columns)}")
    for row in last_table.rows:
        print(f"  {[c.text for c in row.cells]}")

    doc.save(TEMPLATE_PATH)
    print(f"Template saved to {TEMPLATE_PATH}")


if __name__ == '__main__':
    main()
